import os
import uuid
import json
import logging
from datetime import datetime
from flask import Flask, render_template, request, jsonify, redirect, url_for, send_file, send_from_directory, abort
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from functools import wraps
import google.generativeai as genai
from docx import Document
from docxtpl import DocxTemplate
import pandas as pd
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.drawing.spreadsheet_drawing import OneCellAnchor, AnchorMarker
from openpyxl.drawing.xdr import XDRPositiveSize2D
from openpyxl.utils.units import pixels_to_EMU
from deepdiff import DeepDiff

# Local imports
try:
    import database
except ImportError:
    from . import database

# Load environment variables
load_dotenv()

# Configure Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize App
app = Flask(__name__)
app.secret_key = os.getenv('SECRET_KEY', 'default-dev-secret-key')
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['ALLOWED_EXTENSIONS'] = {'docx', 'xlsx', 'png', 'jpg', 'jpeg'}
app.config['MAX_CONTENT_LENGTH'] = 128 * 1024 * 1024  # 128MB

# Ensure upload directory
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Configure Gemini
api_key = os.getenv('GEMINI_API_KEY')
if api_key:
    genai.configure(api_key=api_key)
else:
    logger.warning("GEMINI_API_KEY not found in environment variables.")

# Setup Login Manager
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

class User(UserMixin):
    def __init__(self, id, name, role='operator'):
        self.id = id
        self.name = name
        self.role = role

@login_manager.user_loader
def load_user(user_id):
    user_data = database.get_user(user_id)
    if user_data:
        return User(id=user_id, name=user_data['name'], role=user_data.get('role', 'operator'))
    return None

def role_required(roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for('login'))
            
            # Developer has access to everything
            if current_user.role == 'developer':
                return f(*args, **kwargs)
                
            if current_user.role not in roles:
                return render_template('index.html', error="權限不足：您沒有權限執行此操作。"), 403
            
            return f(*args, **kwargs)
        return decorated_function
    return decorator

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def save_uploaded_file(file):
    if file and allowed_file(file.filename):
        # 4.1 File Handling Logic
        original_filename = file.filename
        file_ext = os.path.splitext(original_filename)[1]
        safe_filename = f"{uuid.uuid4()}{file_ext}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        file.save(filepath)
        return filepath, original_filename
    return None, None

def extract_docx_structure(filepath, limit=2000):
    """提取 Word 文件的結構化資訊 (段落與表格) [New for DeepDiff]"""
    if not filepath or not os.path.exists(filepath):
        return {}
        
    try:
        doc = Document(filepath)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # 提取段落 (含索引，方便比對位置)
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text: # 忽略完全空白行
                structure["paragraphs"].append({
                    "index": i,
                    "text": text[:500] 
                })
                
        # 提取表格 (含坐標)
        for t_idx, table in enumerate(doc.tables):
            table_data = []
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        table_data.append({
                            "loc": f"T{t_idx}:R{r_idx}:C{c_idx}",
                            "text": text
                        })
            if table_data:
                structure["tables"].append(table_data)
                
        return structure
    except Exception as e:
        logger.error(f"Docx structure error: {e}")
        return {}

def extract_xlsx_structure(filepath, limit_rows=100):
    """提取 Excel 文件的結構化資訊 (Sheet 與 Cell) [New for DeepDiff]"""
    if not filepath or not os.path.exists(filepath):
        return {}

    try:
        # Load without read_only to access images. data_only=True for Formula values.
        wb = openpyxl.load_workbook(filepath, data_only=True)
        structure = {
            "sheet_names": wb.sheetnames,
            "sheets": {}
        }
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            cells_data = {}
            row_count = 0
            
            # 1. Extract Text Data
            for row in ws.iter_rows():
                if row_count > limit_rows: break
                for cell in row:
                    if cell.value is not None:
                        # 記錄坐標與值
                        cells_data[f"{cell.row},{cell.column}"] = str(cell.value).strip()
                row_count += 1
            
            # 2. Extract Images (Marker only)
            # This is crucial for "Photo Evaluation"
            try:
                # openpyxl 3.0+ uses ws._images or ws.images
                images_list = getattr(ws, '_images', []) or getattr(ws, 'images', [])
                for img in images_list:
                    # Attempt to find anchor (Top-Left)
                    r, c = None, None
                    
                    # Handling different anchor types (OneCell, TwoCell, Absolute)
                    anchor = img.anchor
                    if hasattr(anchor, '_from'): # TwoCellAnchor
                        r = anchor._from.row + 1 # 0-index to 1-index
                        c = anchor._from.col + 1
                    elif hasattr(anchor, 'row'): # OneCellAnchor (sometimes)
                        r = anchor.row + 1
                        c = anchor.col + 1
                        
                    if r and c:
                        key = f"{r},{c}"
                        exist_val = cells_data.get(key, "")
                        
                        # Extract dimensions avoiding distoration
                        w = getattr(img, 'width', 0)
                        h = getattr(img, 'height', 0)
                        marker = f"<<IMAGE_PRESENT|W:{w}|H:{h}>>"
                        
                        cells_data[key] = f"{exist_val} {marker}".strip()
                        
            except Exception as img_err:
                logger.warning(f"Excel Image extraction warning: {img_err}")
            
            structure["sheets"][sheet_name] = {
                "cells": cells_data
            }
        return structure
    except Exception as e:
        logger.error(f"Excel structure error: {e}")
        return {}

def extract_docx_text(filepath):
    """Legacy extractor (kept for fallback)"""
    try:
        doc = Document(filepath)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return ""

def extract_excel_text(filepath):
    """Legacy extractor (kept for fallback)"""
    try:
        df = pd.read_excel(filepath)
        # Convert first few rows to string representation
        return df.head(10).to_string()
    except Exception as e:
        logger.error(f"Error reading Excel: {e}")
        return ""

# --- Routes ---

@app.route('/')
@login_required
def index():
    projects = database.get_all_projects()
    return render_template('index.html', projects=projects, user=current_user)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        user_id = request.form.get('username')
        password = request.form.get('password')
        user = database.verify_user(user_id, password)
        if user:
            login_user(User(id=user['id'], name=user['name'], role=user.get('role', 'operator')))
            return redirect(url_for('index'))
        else:
            return render_template('login.html', error="Invalid credentials")
            
    return render_template('login.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/developer')
@login_required
@role_required(['developer'])
def developer_dashboard():
    return render_template('developer_dashboard.html', user=current_user)

@app.route('/api/admin/config', methods=['GET', 'POST'])
@login_required
@role_required(['developer'])
def api_admin_config():
    if request.method == 'GET':
        return jsonify(database.get_system_config())
    else:
        new_config = request.json
        current = database.get_system_config()
        current.update(new_config)
        database.save_system_config(current)
        return jsonify({'success': True})

@app.route('/api/admin/stats')
@login_required
@role_required(['developer'])
def api_admin_stats():
    return jsonify(database.get_token_usage_stats())

@app.route('/api/admin/projects')
@login_required
@role_required(['developer'])
def api_admin_projects():
    return jsonify(database.get_all_projects())

@app.route('/api/admin/project/<project_id>', methods=['DELETE'])
@login_required
@role_required(['developer'])
def api_admin_delete_project(project_id):
    success = database.delete_project(project_id)
    if success:
        return jsonify({'success': True})
    return jsonify({'error': 'Delete failed'}), 500

@app.route('/project_setup')
@login_required
@role_required(['manager'])
def project_setup():
    return render_template('project_setup.html')

@app.route('/api/upload', methods=['POST'])
@login_required
def api_upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    if file and allowed_file(file.filename):
        filepath, original_filename = save_uploaded_file(file)
        if filepath:
            filename = os.path.basename(filepath)
            file_size = os.path.getsize(filepath)
            return jsonify({
                'success': True, 
                'file_id': filename, 
                'original_name': original_filename,
                'size': file_size
            })
    return jsonify({'error': 'File type not allowed'}), 400

@app.route('/api/analyze', methods=['POST'])
@role_required(['manager'])
@login_required
def api_analyze():
    data = request.json
    logger.info(f"Analyze request received. Data keys: {data.keys()}")
    
    template_file_id = data.get('template_file_id')
    excel_file_id = data.get('excel_file_id') # Optional
    old_doc_file_id = data.get('old_doc_file_id') # Optional
    
    logger.info(f"Template: {template_file_id}, OldDoc: {old_doc_file_id}")

    if not template_file_id:
        return jsonify({'error': 'Missing template file'}), 400
        
    # Check for same file
    if template_file_id == old_doc_file_id:
        return jsonify({
            'parameters': [],
            'diff_report': '錯誤：您上傳了完全相同的檔案 (ID 相同)。請確保 Zone A (範本) 與 Zone B (舊範例) 是不同的文件。',
            'warning': 'Duplicate files detected.'
        })

    # Paths
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file_id)

    # === [Phase 1 & 2] Structure Extraction & DeepDiff ===
    
    # 1. Extract Template Structure
    ext = os.path.splitext(template_path)[1].lower()
    template_type = "word" if ext in ['.docx', '.doc'] else "excel"
    
    if template_type == "word":
        blank_structure = extract_docx_structure(template_path)
    else:
        blank_structure = extract_xlsx_structure(template_path)
    
    # [Debug] Log Structure Size
    logger.info(f"Blank structure keys: {list(blank_structure.keys()) if blank_structure else 'Empty'}")
    if template_type == "excel" and "sheets" in blank_structure:
        for sname, sdata in blank_structure["sheets"].items():
             logger.info(f"Sheet '{sname}' cells count: {len(sdata.get('cells', {}))}")

    # 2. Extract Reference Structure & Calculate Diff
    formatted_diff_report = "無參考範例，將進行純靜態分析。"
    filled_structure = {}
    
    if old_doc_file_id:
        old_doc_path = os.path.join(app.config['UPLOAD_FOLDER'], old_doc_file_id)
        ref_ext = os.path.splitext(old_doc_path)[1].lower()
        
        # Simple safeguard: compare same file types
        is_same_type = (template_type == "word" and ref_ext in ['.docx', '.doc']) or \
                       (template_type == "excel" and ref_ext in ['.xlsx', '.xls'])
                       
        if is_same_type:
            try:
                if template_type == "word":
                    filled_structure = extract_docx_structure(old_doc_path)
                else:
                    filled_structure = extract_xlsx_structure(old_doc_path)
                
                logger.info(f"Filled structure extracted.")

                # --- DeepDiff Core ---
                # ignore_order=False ensures exact positional matching, vital for forms
                diff = DeepDiff(blank_structure, filled_structure, ignore_order=False, view='tree')
                changes = []
                
                # Value Changes (The most common filling action)
                if 'values_changed' in diff:
                    for node in diff['values_changed']:
                        # path example: ['sheets', 'Sheet1', 'cells', '2,2']
                        path = " -> ".join([str(k) for k in node.path(output_format='list')])
                        old_val = node.t1
                        new_val = node.t2
                        changes.append(f"[內容變更] 位置: {path} | 原始: '{old_val}' -> 填寫: '{new_val}'")

                # Item Added (e.g., repeating rows in tables)
                if 'dictionary_item_added' in diff:
                    for node in diff['dictionary_item_added']:
                        path = " -> ".join([str(k) for k in node.path(output_format='list')])
                        val = node.t2
                        changes.append(f"[新增內容] 位置: {path} | 內容: '{val}'")
                
                # [Fallback Strategy] If DeepDiff finds nothing, try Direct Text Comparison for AI
                if not changes and template_type == "excel":
                     logger.warning("DeepDiff found ZERO changes. Attempting fallback to AI structural inference.")
                     changes.append("[系統備註] DeepDiff 未發現結構差異。使用原始 Excel 內容請求 AI 進行視覺化推斷。")
                     # We will feed the raw structure to AI later if changes is empty, but let's flag it here.
                
                if changes:
                    logger.info(f"DeepDiff found {len(changes)} changes.")
                    formatted_diff_report = "\n".join(changes)
                else:
                    logger.warning("DeepDiff found ZERO changes.")
                    formatted_diff_report = "警告：程式比對後未發現顯著結構差異。\n可能原因：\n1. 兩份文件內容可能完全一致。\n2. 圖片浮動於儲存格上方未被錨定。\n3. 使用了特殊排版(如純文字方塊)導致無法讀取。"

            except Exception as e:
                logger.error(f"DeepDiff/Structure processing failed: {e}")
                formatted_diff_report = f"結構比對失敗: {e}"
        else:
            logger.warning("File types mismatch for diff.")
            formatted_diff_report = "警告：範例檔案格式與模板不符，跳過結構比對。"

    logger.info(f"Diff report length: {len(formatted_diff_report)}")
    # 3. Extract Excel Data Headers (for naming reference)
    excel_headers = []
    if excel_file_id:
        try:
            excel_path = os.path.join(app.config['UPLOAD_FOLDER'], excel_file_id)
            df = pd.read_excel(excel_path)
            excel_headers = df.columns.tolist()
        except Exception as e:
            logger.error(f"Excel header extraction failed: {e}")


    # === [Phase 3] AI Logic Inference ===
    try:
        # Load System Config
        system_config = database.get_system_config()
        model_name = system_config.get('model_name', 'gemini-3-flash-preview')
        prompt_template = system_config.get('ai_prompt_template', database.DEFAULT_SYSTEM_PROMPT)

        logger.info(f"Initializing AI Model: {model_name}")
        model = genai.GenerativeModel(model_name)
        
        # Prepare context data (truncate to avoid token limit)
        blank_json = json.dumps(blank_structure, ensure_ascii=False)[:30000]
        filled_json = json.dumps(filled_structure, ensure_ascii=False)[:30000] if filled_structure else "{}"

        # Analyze Sheet Differences for Logic
        template_sheets = blank_structure.get("sheet_names", [])
        filled_sheets = filled_structure.get("sheet_names", []) if filled_structure else []
        new_sheets = [s for s in filled_sheets if s not in template_sheets]

        # Format the prompt dynamically
        # Consider using safe formatting to avoid KeyError if template has extra braces, 
        # but here we assume the template is controlled.
        # We need to escape brace characters in JSON strings for f-string? 
        # Actually simplest is to NOT use f-string for the template content, but use string.format()
        
        prompt = prompt_template.format(
            template_type=template_type,
            formatted_diff_report=formatted_diff_report[:30000],
            template_sheets=template_sheets,
            filled_sheets=filled_sheets,
            new_sheets=new_sheets,
            blank_json=blank_json,
            filled_json=filled_json
        )

        response = model.generate_content(prompt)
        # Clean response text if it contains markdown code blocks
        text_resp = response.text.replace('```json', '').replace('```', '').strip()
        result_json = json.loads(text_resp)
        
        # Add Diff Report & Token Usage for Step 2 UI
        result_json['diff_report'] = formatted_diff_report

        # Extract Token Usage & Log it
        if response.usage_metadata:
             tokens = {
                 'prompt_tokens': response.usage_metadata.prompt_token_count,
                 'candidates_tokens': response.usage_metadata.candidates_token_count,
                 'total_tokens': response.usage_metadata.total_token_count
             }
             result_json['token_usage'] = tokens
             # Log to database for Developer Dashboard
             database.log_token_usage('unknown_analysis_stage', tokens)
        
        return jsonify(result_json)
        
    except Exception as e:
        logger.error(f"AI Analysis failed: {e}")
        return jsonify({
            'parameters': [],
            'warning': f"AI Analysis Error: {str(e)}"
        })

def create_template(source_path, params):
    """
    Convert original Docx/Excel to Template by replacing original_text with {{ tags }}
    """
    ext = os.path.splitext(source_path)[1].lower()
    
    try:
        if ext in ['.docx', '.doc']:
            doc = Document(source_path)
            for p in params:
                target = p.get('original_text')
                var_name = p.get('name')
                
                if target and var_name:
                    tag = f"{{{{ {var_name} }}}}"
                    
                    # Replace in Paragraphs
                    for para in doc.paragraphs:
                        if target in para.text:
                            para.text = para.text.replace(target, tag)
                    
                    # Replace in Tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if target in cell.text:
                                    cell.text = cell.text.replace(target, tag)
                                    
            new_filename = f"Template_{uuid.uuid4()}.docx"
            new_path = os.path.join(app.config['UPLOAD_FOLDER'], new_filename)
            doc.save(new_path)
            return new_filename

        elif ext in ['.xlsx', '.xls']:
            wb = openpyxl.load_workbook(source_path)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            for p in params:
                                target = p.get('original_text')
                                var_name = p.get('name')
                                if target and var_name and target in cell.value:
                                    tag = f"{{{{ {var_name} }}}}"
                                    cell.value = cell.value.replace(target, tag)

            new_filename = f"Template_{uuid.uuid4()}.xlsx"
            new_path = os.path.join(app.config['UPLOAD_FOLDER'], new_filename)
            wb.save(new_path)
            return new_filename

    except Exception as e:
        logger.error(f"Template conversion failed: {e}")
        return None

@app.route('/api/save_project', methods=['POST'])
@role_required(['manager'])
@login_required
def api_save_project():
    data = request.json
    project_id = str(uuid.uuid4())
    
    # Create the template
    template_file_id = data.get('template_file_id')
    parameters = data.get('parameters', [])
    
    final_template_name = template_file_id
    if template_file_id:
        source_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file_id)
        if os.path.exists(source_path):
            converted_name = create_template(source_path, parameters)
            if converted_name:
                final_template_name = converted_name
    
    config = {
        'name': data.get('project_name', 'System Project'),
        'description': data.get('project_desc', ''),
        'mode': data.get('mode', 'one_shot'),
        'features': data.get('features', {}),
        'created_at': datetime.now().isoformat(),
        'template_file': final_template_name,
        'parameters': parameters
    }
    
    database.save_project_config(project_id, config)
    return jsonify({'success': True, 'project_id': project_id})


@app.route('/project/<project_id>')
@login_required
def project_form(project_id):
    config = database.get_project_config(project_id)
    if not config:
        return redirect(url_for('index'))
    
    # Route based on Mode
    if config.get('mode') == 'monthly':
        return render_template('project_dashboard_monthly.html', project=config, project_id=project_id)
        
    return render_template('project_form.html', project=config, project_id=project_id)

@app.route('/generate_document', methods=['POST'])
@login_required
def generate_document():
    project_id = request.form.get('project_id')
    config = database.get_project_config(project_id)
    if not config:
        return "Project not found", 404
        
    # Get form data mapping
    context = {}
    for param in config['parameters']:
        key = param['name']
        val = request.form.get(key, '')
        context[key] = val
        
    # Generation Logic
    template_file = config['template_file']
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file)
    ext = os.path.splitext(template_file)[1].lower()
    
    if not os.path.exists(template_path):
        return "Template file missing", 404

    try:
        output_filename = f"Generated_{config.get('name', 'Doc')}_{datetime.now().strftime('%Y%m%d%H%M')}{ext}"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        if ext == '.docx':
            doc = DocxTemplate(template_path)
            doc.render(context)
            doc.save(output_path)
            
        elif ext in ['.xlsx', '.xls']:
            wb = openpyxl.load_workbook(template_path)
            # Find and replace {{ key }} with val
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            for key, val in context.items():
                                tag = f"{{{{ {key} }}}}"
                                if tag in cell.value:
                                    # If exact match, perform type adjustment if possible (number)
                                    if cell.value.strip() == tag:
                                        # Try to convert to number if possible
                                        try:
                                            cell.value = float(val) if '.' in val else int(val)
                                        except ValueError:
                                            cell.value = val
                                    else:
                                        cell.value = cell.value.replace(tag, str(val))
            wb.save(output_path)
        
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        logger.error(f"Generation failed: {e}")
        return f"Error generating document: {e}", 500

# --- Monthly Project APIs ---

@app.route('/api/project/<project_id>/entries', methods=['GET'])
@login_required
def api_get_entries(project_id):
    entries = database.get_project_entries(project_id)
    return jsonify({'entries': entries})

@app.route('/api/project/<project_id>/entry', methods=['POST'])
@login_required
def api_add_entry(project_id):
    try:
        # Handle form data and files
        config = database.get_project_config(project_id)
        if not config:
            return jsonify({'success': False, 'error': 'Project not found'})
            
        logger.info(f"Adding entry for project {config.get('name')}. Found {len(config.get('parameters', []))} parameters.")
        
        # DEBUG: Log incoming request data
        print(f"DEBUG: Form Keys: {list(request.form.keys())}")
        print(f"DEBUG: File Keys: {list(request.files.keys())}")

        entry_data = {
            'id': str(uuid.uuid4()),
            'date': request.form.get('entry_date'),
            'created_at': datetime.now().isoformat(),
            'data': {}
        }
        
        # Process parameters
        for param in config.get('parameters', []):
            field = param['name']
            # logger.info(f"Processing field: {field}, Type: {param['type']}")
            
            if param['type'] == 'image':
                file = request.files.get(field)
                print(f"DEBUG: Processing Image '{field}' - File Object: {file}")
                if file:
                    filepath, orig_name = save_uploaded_file(file)
                    # We store the relative path for the frontend/processing
                    if filepath:
                        # Make path relative to uploads/ for serving
                        rel_path = os.path.relpath(filepath, app.config['UPLOAD_FOLDER'])
                        entry_data['data'][field] = f"uploads/{rel_path}"
                        logger.info(f"Saved image for {field}: {entry_data['data'][field]}")
                else:
                    entry_data['data'][field] = None
                    print(f"DEBUG: No file found for '{field}'")
            else:
                entry_data['data'][field] = request.form.get(field, '')
                
        database.save_project_entry(project_id, entry_data)
        return jsonify({'success': True})
        
    except Exception as e:
        logger.error(f"Error adding entry: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/project/<project_id>/entry/<entry_id>', methods=['DELETE'])
@login_required
def api_delete_entry(project_id, entry_id):
    try:
        database.delete_project_entry(project_id, entry_id)
        return jsonify({'success': True})
    except Exception as e:
        logger.error(f"Error deleting entry: {e}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/api/project/<project_id>/generate_monthly', methods=['POST'])
@login_required
def api_generate_monthly(project_id):
    config = database.get_project_config(project_id)
    if not config:
        return "Project not found", 404
        
    entries = database.get_project_entries(project_id)
    if not entries:
         return "No entries to generate", 400

    template_file = config['template_file']
    template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file)
    
    if not os.path.exists(template_path):
        return "Template file missing", 404
        
    try:
        wb = openpyxl.load_workbook(template_path)
        # Assume the first sheet is the template to copy
        source_sheet = wb.worksheets[0]
        
        # Sort entries by date
        entries.sort(key=lambda x: x['date'])
        
        for entry in entries:
            target_sheet = wb.copy_worksheet(source_sheet)
            
            # 1. Determine Sheet Name
            sheet_name = entry['date'][5:] # Default MMDD from YYYY-MM-DD
            
            # Try to find specific sheet_name param
            for param in config.get('parameters', []):
                 if param['name'] == 'sheet_name' and entry['data'].get('sheet_name'):
                      raw_name = entry['data']['sheet_name']
                      if raw_name: sheet_name = raw_name
                      break
            
            # Sanitize sheet name
            safe_name = str(sheet_name).replace(':', '').replace('/', '-').replace('\\', '').replace('?', '').replace('*', '').replace('[', '').replace(']', '')[:30]
            target_sheet.title = safe_name
            
            # 2. Fill Data & Images
            data_map = entry['data']
            
            for param in config.get('parameters', []):
                key = param['name']
                val = data_map.get(key)
                original_text = param.get('original_text', '')
                
                # Tag to look for in cells (e.g. {{ removal_photo }})
                tag = f"{{{{ {key} }}}}"
                
                if param['type'] == 'image':
                    # --- Image Logic ---
                    # We always search for the tag to handle text restoration or clearing
                    # But we also use anchor for insertion
                    
                    anchor = param.get('style', {}).get('anchor_cell')
                    image_inserted = False
                    
                    if val:
                        # Try to insert image
                        try:
                            # Handle different path formats
                            # Check if val is an absolute path that exists (for local testing/manual entry)
                            if os.path.isabs(val) and os.path.exists(val):
                                img_path = val
                            else:
                                # Default: look in UPLOAD_FOLDER
                                filename = os.path.basename(val)
                                img_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                            
                            if os.path.exists(img_path):
                                img = OpenpyxlImage(img_path)
                                
                                # Smart Positioning & Resizing
                                final_anchor = "A1"
                                is_header_anchor = False
                                
                                if anchor and ',' in str(anchor):
                                    r_start, c_start = map(int, anchor.split(','))
                                    final_anchor = f"{get_column_letter(c_start)}{r_start}"
                                    
                                    # Find if anchor is in a merged range
                                    anchor_range = None
                                    for rng in target_sheet.merged_cells.ranges:
                                        if rng.min_row <= r_start <= rng.max_row and rng.min_col <= c_start <= rng.max_col:
                                            anchor_range = rng
                                            break
                                    
                                    target_box = anchor_range
                                    
                                    # Check if it's a "Header" (small height)
                                    if anchor_range:
                                        total_h = 0
                                        for ri in range(anchor_range.min_row, anchor_range.max_row + 1):
                                            rh = target_sheet.row_dimensions[ri].height
                                            total_h += rh if rh else 15
                                        
                                        # Threshold: 60 points (approx 80px)
                                        if total_h < 60:
                                            # It is likely a header. Look for the BODY below it.
                                            is_header_anchor = True
                                            next_row = anchor_range.max_row + 1
                                            
                                            # Find merged range starting at next_row, same col
                                            body_range = None
                                            for rng in target_sheet.merged_cells.ranges:
                                                if rng.min_row == next_row and rng.min_col == c_start:
                                                    body_range = rng
                                                    break
                                            
                                            if body_range:
                                                target_box = body_range
                                                final_anchor = f"{get_column_letter(c_start)}{next_row}"
                                    
                                    # Calculate Target Dimensions (Pixels)
                                    if target_box:
                                        # Width
                                        box_w = 0
                                        for ci in range(target_box.min_col, target_box.max_col + 1):
                                            cw = target_sheet.column_dimensions[get_column_letter(ci)].width
                                            box_w += (cw if cw else 8.43) * 7.5 # Approx 7.5 px per unit
                                            
                                        # Height
                                        box_h = 0
                                        for ri in range(target_box.min_row, target_box.max_row + 1):
                                            rh = target_sheet.row_dimensions[ri].height
                                            box_h += (rh if rh else 15) * 1.333 # Approx 1.33 px per point
                                            
                                        # Resize Image (Aspect Fit with padding)
                                        if box_w > 10 and box_h > 10:
                                            img_ratio = img.width / img.height
                                            box_ratio = box_w / box_h
                                            
                                            # Padding
                                            pad_w = min(20, box_w * 0.1)
                                            pad_h = min(20, box_h * 0.1)
                                            avail_w = box_w - pad_w
                                            avail_h = box_h - pad_h
                                            
                                            if img_ratio > box_ratio:
                                                # Fit to Width
                                                new_w = avail_w
                                                new_h = avail_w / img_ratio
                                            else:
                                                # Fit to Height
                                                new_h = avail_h
                                                new_w = avail_h * img_ratio
                                                
                                            img.width = new_w
                                            img.height = new_h

                                # Calculate Offsets for Centering
                                # Note: AnchorMarker uses 0-indexed row/col
                                col_idx_0 = (c_start - 1)
                                row_idx_0 = (next_row - 1 if is_header_anchor else r_start - 1)
                                
                                # Use box dims vs new image dims to find center
                                # box_w, box_h are in "approximate pixels" (from our calculation loop)
                                # But Openpyxl pixels_to_EMU is exact (1px = 9525 EMU)
                                # Our box_w calculation: (cw * 7.5) -> approximate
                                # It's better to center based on the available space we calculated
                                
                                offset_x_px = (box_w - new_w) / 2
                                offset_y_px = (box_h - new_h) / 2
                                
                                # Ensure non-negative
                                offset_x_px = max(0, offset_x_px)
                                offset_y_px = max(0, offset_y_px)
                                
                                offset_x_emu = int(pixels_to_EMU(offset_x_px))
                                offset_y_emu = int(pixels_to_EMU(offset_y_px))
                                
                                # Create proper OneCellAnchor to support offsets
                                marker = AnchorMarker(col=col_idx_0, colOff=offset_x_emu, row=row_idx_0, rowOff=offset_y_emu)
                                size_emu = XDRPositiveSize2D(int(pixels_to_EMU(new_w)), int(pixels_to_EMU(new_h)))
                                
                                img.anchor = OneCellAnchor(_from=marker, ext=size_emu)
                                target_sheet.add_image(img)
                                image_inserted = True
                            else:
                                logger.warning(f"Image not found: {img_path}")
                                
                        except Exception as img_err:
                            logger.error(f"Image insert error {key}: {img_err}")
                    
                    # --- Text Cleanup for Image Placeholder ---
                    # If is_header_anchor is True -> We moved the image down, so RESTORE the original header text
                    # Otherwise -> Clear the text
                    
                    if image_inserted:
                        replacement_text = original_text if is_header_anchor else ""
                    else:
                        replacement_text = original_text
                    
                    # Note: Anchor cell clearing was problematic if anchor != placeholder cell
                    # So we scan for the placeholder TAG explicitly.
                    for row in target_sheet.iter_rows():
                        for cell in row:
                            if cell.value and isinstance(cell.value, str) and tag in cell.value:
                                if cell.value.strip() == tag:
                                    cell.value = replacement_text
                                else:
                                    cell.value = cell.value.replace(tag, replacement_text)

                else:
                    # --- Text Logic ---
                    if val is None: val = "" # Handle None
                    
                    # Search and replace
                    for row in target_sheet.iter_rows():
                        for cell in row:
                             if cell.value and isinstance(cell.value, str) and tag in cell.value:
                                 # Exact or partial match
                                 if cell.value.strip() == tag:
                                     # Try number conversion
                                     try:
                                         if str(val).isdigit():
                                             cell.value = int(val)
                                         elif str(val).replace('.', '', 1).isdigit():
                                              cell.value = float(val)
                                         else:
                                              cell.value = val
                                     except:
                                         cell.value = val
                                 else:
                                     cell.value = cell.value.replace(tag, str(val))
        
        # Remove the original template sheet
        if len(wb.sheetnames) > 1:
            wb.remove(source_sheet)
        
        # Save
        if config.get('name'):
             out_name = f"{config['name']}_月報_{datetime.now().strftime('%Y%m')}.xlsx"
        else:
             out_name = f"Monthly_Report_{project_id}.xlsx"
             
        out_path = os.path.join(app.config['UPLOAD_FOLDER'], out_name)
        wb.save(out_path)
        
        return send_file(out_path, as_attachment=True, download_name=out_name)
        
    except Exception as e:
        logger.error(f"Monthly Generation Failed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    # Ensure server reloads if this file changes...
    app.run(debug=True, port=5000)
